"""
Generate a Project Closure Certificate (.docx) for a student using the
template file  'project closing certificate.docx'.

Usage (from the web app):
    python generate_closure.py --student-json '{"name":"...","start_date":"...","end_date":"..."}'

The certificate is saved into the 'output' folder.
"""

import os
import sys
import json
import argparse
import datetime
import tempfile
import time

from docx import Document
from docx.shared import Pt

TEMPLATE_FILE = "../sample/project closing certificate.docx"
OUTPUT_FOLDER = "output"


VIRTUAL_PRINTERS = {"microsoft print to pdf", "microsoft xps document writer",
                     "onenote", "onenote (desktop)", "onenote (desktop) - protected",
                     "fax", "send to onenote"}


def find_physical_printer():
    """
    Return the printer to use.
    Priority:
      1. NRSC_PRINTER environment variable (set this to the exact LAN printer name)
      2. First non-virtual printer found on the system
    """
    env_printer = os.environ.get("NRSC_PRINTER", "").strip()
    if env_printer:
        return env_printer
    import subprocess
    result = subprocess.run(
        ["powershell", "-NoProfile", "-Command",
         "Get-Printer | Select-Object -ExpandProperty Name"],
        capture_output=True, text=True
    )
    for line in result.stdout.strip().splitlines():
        name = line.strip()
        if name and name.lower() not in VIRTUAL_PRINTERS:
            return name
    return None


def print_docx(filepath):
    """Open .docx in Word visibly for review — user can edit and print manually."""
    import pythoncom
    import win32com.client

    abs_path = os.path.abspath(filepath)
    pythoncom.CoInitialize()
    try:
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = True
        word.WindowState = 1            # wdWindowStateNormal
        word.Documents.Open(abs_path)
        # Leave Word open — user reviews/prints manually
    except Exception:
        raise
    finally:
        pythoncom.CoUninitialize()


# ─────────────────────────────────────────────────────────────────────────────
# Helpers
# ─────────────────────────────────────────────────────────────────────────────

def format_date(date_str: str) -> str:
    """Parse an ISO-ish date string and return  DD Mon YYYY  (e.g. 01 Feb 2026)."""
    if not date_str:
        return ""
    try:
        dt = datetime.datetime.fromisoformat(date_str.replace(" ", "T").split("T")[0])
        return dt.strftime("%d %b %Y")
    except ValueError:
        return date_str


def build_replacements(student: dict) -> dict:
    """Return placeholder -> value mapping for the closure certificate template."""
    name       = student.get("name", "").strip()
    start_date = format_date(student.get("start_date", ""))
    end_date   = format_date(student.get("end_date", ""))

    return {
        "{{End_Date}}":      end_date,
        "{{Start_Date}}":    start_date,
        "{{Student_Name}}":  name,
        "{{Present_Date}}":  datetime.datetime.now().strftime("%d %b %Y"),
    }


def replace_in_paragraph(para, replacements):
    """
    Replace every placeholder in *para*, including those split across runs.

    Pass 1 -- per-run:  fast path for placeholders entirely within one run.
    Pass 2 -- merged :  concatenate all run texts, replace ALL remaining
              placeholders (handles brace-style {{...}} that Word splits
              across multiple runs), then put the result back into the
              first run and blank the rest.
    """
    # Pass 1: replace within individual runs (preserves per-run formatting)
    for run in para.runs:
        for ph, val in replacements.items():
            if ph in run.text:
                run.text = run.text.replace(ph, val)

    # Pass 2: handle any placeholders still spanning multiple runs
    full = "".join(r.text for r in para.runs)
    if not any(ph in full for ph in replacements):
        return

    for ph, val in replacements.items():
        full = full.replace(ph, val)

    if para.runs:
        para.runs[0].text = full
        for run in para.runs[1:]:
            run.text = ""


def replace_all(doc, replacements):
    """Apply replacements to every paragraph, table cell, header, and footer."""
    for para in doc.paragraphs:
        replace_in_paragraph(para, replacements)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    replace_in_paragraph(para, replacements)

    for section in doc.sections:
        for para in section.header.paragraphs:
            replace_in_paragraph(para, replacements)
        for para in section.footer.paragraphs:
            replace_in_paragraph(para, replacements)


# ─────────────────────────────────────────────────────────────────────────────
# Certificate generation
# ─────────────────────────────────────────────────────────────────────────────

def _bold_fragments(para, fragments):
    """
    Split para runs so that only the text in *fragments* is bold,
    leaving the surrounding text with its original bold state.
    """
    import copy
    from docx.oxml.ns import qn

    # Collapse all runs into a single run (text already merged by replace_all)
    full_text = "".join(r.text for r in para.runs)
    if not full_text:
        return

    # Build (segment_text, is_bold) pairs by scanning full_text
    segments = []
    remaining = full_text
    while remaining:
        earliest_pos = len(remaining)
        earliest_frag = None
        for frag in fragments:
            pos = remaining.find(frag)
            if pos != -1 and pos < earliest_pos:
                earliest_pos = pos
                earliest_frag = frag
        if earliest_frag is None:
            segments.append((remaining, False))
            break
        if earliest_pos > 0:
            segments.append((remaining[:earliest_pos], False))
        segments.append((earliest_frag, True))
        remaining = remaining[earliest_pos + len(earliest_frag):]

    # Clear existing runs
    p = para._p
    existing_runs = p.findall(qn('w:r'))
    template_rpr = None
    if existing_runs:
        rpr = existing_runs[0].find(qn('w:rPr'))
        if rpr is not None:
            template_rpr = copy.deepcopy(rpr)
        for r in existing_runs:
            p.remove(r)

    # Re-create runs with correct bold state
    for text, is_bold in segments:
        if not text:
            continue
        run = para.add_run(text)
        run.bold = is_bold
        # Preserve original run properties (font, size etc.) if available
        if template_rpr is not None:
            new_rpr = copy.deepcopy(template_rpr)
            # Override bold element
            for b_tag in (qn('w:b'), qn('w:bCs')):
                for old in new_rpr.findall(b_tag):
                    new_rpr.remove(old)
            from lxml import etree
            if is_bold:
                new_rpr.append(etree.SubElement(new_rpr, qn('w:b')))
                new_rpr.append(etree.SubElement(new_rpr, qn('w:bCs')))
            else:
                # Explicitly suppress bold so paragraph/style bold is not inherited
                b_off = etree.SubElement(new_rpr, qn('w:b'))
                b_off.set(qn('w:val'), '0')
                bcs_off = etree.SubElement(new_rpr, qn('w:bCs'))
                bcs_off.set(qn('w:val'), '0')
            run._r.insert(0, new_rpr)
            # Remove the duplicate rPr that add_run may have added
            rprs = run._r.findall(qn('w:rPr'))
            for extra in rprs[1:]:
                run._r.remove(extra)


def generate_closure_certificate(template_path: str, student: dict, output_folder: str) -> str:
    """Open the template, replace placeholders, bold name/dates, save to output folder."""
    doc          = Document(template_path)
    replacements = build_replacements(student)

    replace_all(doc, replacements)

    # Bold only specific fragments (name and date range), not whole paragraphs
    name       = student.get("name", "").strip()
    start_date = format_date(student.get("start_date", ""))
    end_date   = format_date(student.get("end_date", ""))

    for para in doc.paragraphs:
        txt = para.text
        # Collect which fragments appear in this paragraph
        frags = [f for f in [name, start_date, end_date] if f and f in txt]
        if not frags:
            continue
        # Use fragment-level bolding so only the matched text is bold
        _bold_fragments(para, frags)

    name      = student.get("name", "unknown").strip()
    safe_name = "".join(c if c.isalnum() or c in " _-" else "_" for c in name).strip()
    os.makedirs(output_folder, exist_ok=True)
    out_path  = os.path.join(output_folder, f"Closure_Certificate_{safe_name}.docx")
    doc.save(out_path)
    return out_path


def main():
    parser = argparse.ArgumentParser(description="Generate Project Closure Certificate")
    parser.add_argument("--student-json", required=True, help="JSON object with student data")
    parser.add_argument("--template", default=TEMPLATE_FILE, help="Template .docx file")
    parser.add_argument("--output", default=OUTPUT_FOLDER, help="Output folder path")
    parser.add_argument("--print", action="store_true", dest="send_to_printer", help="Send to default printer instead of saving to output")
    args = parser.parse_args()

    script_dir = os.path.dirname(os.path.abspath(__file__))
    template_path = args.template if os.path.isabs(args.template) else os.path.join(script_dir, args.template)

    output_folder = args.output if os.path.isabs(args.output) else os.path.join(script_dir, args.output)

    if not os.path.exists(template_path):
        print(f"\nTemplate not found: {template_path}")
        print(f"Make sure '{TEMPLATE_FILE}' is in the same folder as this script.")
        sys.exit(1)

    try:
        student = json.loads(args.student_json)
    except json.JSONDecodeError as exc:
        print(f"Invalid --student-json payload: {exc}")
        sys.exit(1)

    if not isinstance(student, dict):
        print("--student-json must be a single JSON object")
        sys.exit(1)

    print("\n" + "=" * 54)
    print("   Project Closure Certificate Generator")
    print("=" * 54)

    try:
        out_path = generate_closure_certificate(template_path, student, output_folder)
        if args.send_to_printer:
            print_docx(out_path)
            print(f"\n  Certificate sent to printer -> {os.path.basename(out_path)}")
            print(f"\n{'='*54}")
            print(f"  1 certificate sent to printer")
            # clean up temp folder
            try:
                os.rmdir(output_folder)
            except OSError:
                pass
        else:
            print(f"\n  Certificate saved -> {os.path.basename(out_path)}")
            print(f"\n{'='*54}")
            print(f"  1 certificate saved to  ->  {output_folder}")
        print("=" * 54 + "\n")
        print(f"RESULT: success=1 failed=0 output={output_folder}")
    except Exception as e:
        print(f"\n  Error: {e}")
        import traceback
        traceback.print_exc()
        print(f"RESULT: success=0 failed=1 output={output_folder}")
        sys.exit(1)


if __name__ == "__main__":
    main()
