import os
import sys
import re
import argparse
import json
import tempfile
import time
import urllib.request
import openpyxl
from docx import Document
from docx.shared import Pt

EXCEL_FILE    = "https://docs.google.com/spreadsheets/d/1xFvqME_wXQTRYkJJl9W-_OT1L_XBUz3RY8gUluhGP8w/edit?usp=sharing"
TEMPLATE_FILE = "../sample/SampleGALetter.docx"
OUTPUT_FOLDER = "output"


VIRTUAL_PRINTERS = {"microsoft print to pdf", "microsoft xps document writer",
                     "onenote", "onenote (desktop)", "onenote (desktop) - protected",
                     "fax", "send to onenote"}


def find_physical_printer():
    """
    Return the printer to use.
    Priority:
      1. NRSC_PRINTER environment variable (set this to the exact LAN printer name)
      2. First non-virtual printer found on the system
    """
    env_printer = os.environ.get("NRSC_PRINTER", "").strip()
    if env_printer:
        return env_printer
    import subprocess
    result = subprocess.run(
        ["powershell", "-NoProfile", "-Command",
         "Get-Printer | Select-Object -ExpandProperty Name"],
        capture_output=True, text=True
    )
    for line in result.stdout.strip().splitlines():
        name = line.strip()
        if name and name.lower() not in VIRTUAL_PRINTERS:
            return name
    return None


def print_docx(filepath):
    """Open .docx in Word visibly for review — user can edit and print manually."""
    import pythoncom
    import win32com.client

    abs_path = os.path.abspath(filepath)
    pythoncom.CoInitialize()
    try:
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = True
        word.WindowState = 1            # wdWindowStateNormal
        word.Documents.Open(abs_path)
        # Leave Word open — user reviews/prints manually
    except Exception:
        raise
    finally:
        pythoncom.CoUninitialize()

def resolve_excel(source):
    """Return a local .xlsx path — downloading from Google Sheets if needed."""
    if not source.startswith("http"):
        return source

    print("Downloading from Google Sheets...")
    match = re.search(r"/spreadsheets/d/([a-zA-Z0-9_-]+)", source)
    if not match:
        print("Could not parse Google Sheets ID from the URL.")
        sys.exit(1)

    sheet_id   = match.group(1)
    export_url = (
        f"https://docs.google.com/spreadsheets/d/{sheet_id}"
        f"/export?format=xlsx"
    )
    try:
        tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".xlsx")
        tmp.close()
        urllib.request.urlretrieve(export_url, tmp.name)
        print("Sheet downloaded.\n")
        return tmp.name
    except Exception as e:
        print(f"Download failed: {e}")
        print("Make sure the sheet is set to  'Anyone with the link can view'.")
        sys.exit(1)


def load_students(source):
    path = resolve_excel(source)
    wb   = openpyxl.load_workbook(path)
    ws   = wb.active
    rows = list(ws.iter_rows(values_only=True))

    if not rows:
        print("The spreadsheet appears to be empty.")
        sys.exit(1)

    headers  = [str(h).strip().lower() if h else "" for h in rows[0]]
    students = []
    for i, row in enumerate(rows[1:], start=2):
        s = {"_row": i}
        for h, v in zip(headers, row):
            s[h] = str(v).strip() if v is not None else ""
        students.append(s)
    return students


# ─────────────────────────────────────────────────────────────────────────────
# Replacement helpers
# ─────────────────────────────────────────────────────────────────────────────

def format_date(date_str):
    """Parse an ISO-ish date string and return  DD Mon YYYY  (e.g. 01 Feb 2026)."""
    import datetime
    if not date_str:
        return ""
    try:
        dt = datetime.datetime.fromisoformat(date_str.replace(" ", "T").split("T")[0])
        return dt.strftime("%d %b %Y")
    except ValueError:
        return date_str


def build_replacements(s):
    """
    Return an ordered dict of placeholder -> value.

    The template uses TWO kinds of placeholders:
      1. Brace-wrapped  {{Placeholder}}  -- always split across three separate
         runs in Word ('{{' / 'Name' / '}}'), handled by the paragraph-merge
         fallback in replace_in_paragraph.
      2. Bare-word  Salute / Course / District / Name_Guide  -- plain text
         inside a single run; handled by the per-run pass.
    """
    salute     = s.get("salute", "")
    name       = s.get("name", "")
    college    = s.get("college name", "")
    district   = s.get("district", "")
    guide_name = s.get("guide name", "")
    guide_area = s.get("guide area", "")
    course     = s.get("course", "")
    from_date  = format_date(s.get("start date", ""))
    to_date    = format_date(s.get("end date", ""))
    issued_dt  = format_date(s.get("allotment date", ""))
    month      = str(s.get("month", "")).strip()
    if not month:
        import datetime
        source_date = str(s.get("allotment date", "") or s.get("start date", "")).strip()
        try:
            dt = datetime.datetime.fromisoformat(source_date.replace(" ", "T").split("T")[0])
            month = dt.strftime("%b")
        except Exception:
            month = ""
    sl_no      = str(s.get("sl. no.", "") or s.get("sl_no", "")).strip()
    guide_ro   = s.get("guide reporting officer", "")
    college_hod = str(s.get("college dean hod", "") or s.get("college_dean_hod", "")).strip()
    dd         = str(s.get("dd", "")).strip()

    return {
        # ── brace-wrapped (split-run) placeholders ──────────────────────────
        "{{Guide_AllotmentLetter_IssuedDate}}": issued_dt,
        "{{Name_Student}}":                    name,
        "{{Salute}}":                          salute,
        "{{Name_College}}":                    college,
        "{{District}}":                        district,
        "{{Name_Guide}}":                      guide_name,
        "{{Guide_Area}}":                      guide_area,
        "{{Course}}":                          course,
        "{{From_Date}}":                       from_date,
        "{{To_Date}}":                         to_date,
        "{{Guide Reporting Officer}}": guide_ro,
        "{{College_Dean_HOD}}": college_hod,
        "{{DD}}": dd,

        # ── bare-word placeholders (no braces in template body) ─────────────
        # "...this is to inform you that Salute {{Name_Student}}..."
        "Salute ":  salute + " ",

        # "...student of Course is permitted..."
        "Course":   course,

        # Standalone 'District' line in the address block
        "{{District}}":   district,

        # "Project Guide - Name_Guide - Kindly ensure..."
        "{{Name_Guide}}": guide_name,

        # NRSC: ORF: 2026/Month/Sl. No.  -- reference number line
        "Month":   month,
        "Sl. No.": sl_no if sl_no else "Sl. No.",
    }


def replace_in_paragraph(para, replacements):
    """
    Replace every placeholder in *para*, including those split across runs.

    Pass 1 -- per-run:  fast path for placeholders entirely within one run
                        (bare-word placeholders: Salute, Course, District,
                         Name_Guide).
    Pass 2 -- merged :  concatenate all run texts, replace ALL remaining
                        placeholders (handles brace-style {{...}} that Word
                        splits across 3 runs), then put the result back into
                        the first run and blank the rest.
                        NOTE: the old code had `break` after the first match
                        which left {{From_Date}}, {{To_Date}}, {{Guide_Area}}
                        unreplaced -- that bug is now fixed.
    """
    # Pass 1: replace within individual runs (preserves per-run formatting)
    for run in para.runs:
        for ph, val in replacements.items():
            if ph in run.text:
                run.text = run.text.replace(ph, val)

    # Pass 2: handle any placeholders still spanning multiple runs
    full = "".join(r.text for r in para.runs)
    if not any(ph in full for ph in replacements):
        return  # nothing left to do

    # Replace ALL remaining placeholders (no break -- fixes the original bug)
    for ph, val in replacements.items():
        full = full.replace(ph, val)

    if para.runs:
        para.runs[0].text = full
        for run in para.runs[1:]:
            run.text = ""


def replace_all(doc, replacements):
    """Apply replacements to every paragraph, table cell, header, and footer."""
    for para in doc.paragraphs:
        replace_in_paragraph(para, replacements)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    replace_in_paragraph(para, replacements)

    for section in doc.sections:
        for para in section.header.paragraphs:
            replace_in_paragraph(para, replacements)
        for para in section.footer.paragraphs:
            replace_in_paragraph(para, replacements)


def bold_fragments_in_paragraph(para, fragments):
    """
    Split runs[0] (which Pass 2 filled with the entire paragraph text) into
    multiple runs, making each text fragment in *fragments* bold.
    """
    from copy import deepcopy
    from lxml import etree

    if not para.runs:
        return

    run0 = para.runs[0]
    full_text = run0.text
    if not full_text:
        return

    # Find where each fragment sits in the full text
    positions = []
    for frag in fragments:
        if frag:
            idx = full_text.find(frag)
            if idx >= 0:
                positions.append((idx, len(frag)))
    if not positions:
        return
    positions.sort()

    # Build (text, is_bold) segments
    segments, ptr = [], 0
    for start, length in positions:
        if start < ptr:
            continue  # overlapping — skip
        if start > ptr:
            segments.append((full_text[ptr:start], False))
        segments.append((full_text[start:start + length], True))
        ptr = start + length
    if ptr < len(full_text):
        segments.append((full_text[ptr:], False))

    if not segments:
        return

    # Save a clean copy of run0's XML to use as template for new runs
    run0_tmpl = deepcopy(run0._r)

    # Apply first segment to run0 in-place
    run0.text = segments[0][0]
    run0.bold = segments[0][1]

    if len(segments) == 1:
        return

    # Insert new <w:r> elements after run0 for the remaining segments
    ns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    p_elem = para._p
    prev_r = run0._r

    for text, is_bold in segments[1:]:
        new_r = deepcopy(run0_tmpl)

        # Set the text content
        w_t = new_r.find(f'{{{ns}}}t')
        if w_t is None:
            w_t = etree.SubElement(new_r, f'{{{ns}}}t')
        w_t.text = text
        # Preserve leading/trailing spaces
        if text and (text[0] == ' ' or text[-1] == ' '):
            w_t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')

        # Ensure rPr block exists
        rPr = new_r.find(f'{{{ns}}}rPr')
        if rPr is None:
            rPr = etree.Element(f'{{{ns}}}rPr')
            new_r.insert(0, rPr)

        # Clear any existing <w:b> and re-set according to is_bold
        for b in rPr.findall(f'{{{ns}}}b'):
            rPr.remove(b)
        if is_bold:
            etree.SubElement(rPr, f'{{{ns}}}b')

        # Insert this new run right after prev_r
        insert_idx = list(p_elem).index(prev_r)
        p_elem.insert(insert_idx + 1, new_r)
        prev_r = new_r


# ─────────────────────────────────────────────────────────────────────────────
# Font-size reduction
# ─────────────────────────────────────────────────────────────────────────────

def adjust_font_sizes(doc, target_pt=10):
    """
    Set every font size in *doc* to exactly *target_pt* pt.

    Covers two storage locations:
      * Styles XML  (docDefaults + named styles) -- sizes stored as half-points
        in  w:sz / w:szCs  attributes.
      * Run-level overrides  (w:rPr/w:sz)  -- exposed as  run.font.size  (EMU).
    """
    from docx.oxml.ns import qn
    half_pts = target_pt * 2      # OOXML half-point units

    # 1. Styles XML (document defaults + all named styles)
    styles_elem = doc.part.styles._element
    for tag in (qn("w:sz"), qn("w:szCs")):
        for elem in styles_elem.iter(tag):
            if elem.get(qn("w:val")) is not None:
                elem.set(qn("w:val"), str(half_pts))

    # 2. Explicit run-level sizes in the document body
    def shrink(run):
        if run.font.size is not None:
            run.font.size = Pt(target_pt)

    for para in doc.paragraphs:
        for run in para.runs:
            shrink(run)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        shrink(run)

    for section in doc.sections:
        for para in section.header.paragraphs:
            for run in para.runs:
                shrink(run)
        for para in section.footer.paragraphs:
            for run in para.runs:
                shrink(run)


# ─────────────────────────────────────────────────────────────────────────────
# Letter generation helpers
# ─────────────────────────────────────────────────────────────────────────────

def _make_para_bold(para):
    """Make every run in *para* bold."""
    for run in para.runs:
        run.bold = True


def _set_para_bold(para, text):
    """Replace all runs in *para* with a single bold run (preserves list numbering)."""
    from docx.oxml.ns import qn
    p = para._p
    for r in p.findall(qn('w:r')):
        p.remove(r)
    run = para.add_run(text)
    run.bold = True


# ─────────────────────────────────────────────────────────────────────────────
# Letter generation
# ─────────────────────────────────────────────────────────────────────────────

def generate_letter(template_path, student, output_folder):
    doc          = Document(template_path)
    replacements = build_replacements(student)

    # Fix HOD line BEFORE replace_all so we can still see the template placeholder.
    # Template has "{{Name_College}} Dean/HOD" in both the To block and Copy-to block.
    # Replace each with just the college_dean_hod value (col 4 of Excel).
    hod_val     = student.get("college dean hod", "").strip()
    college_val = student.get("college name", "").strip()
    def _iter_all_paragraphs(document):
        for paragraph in document.paragraphs:
            yield paragraph
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        yield paragraph
        for section in document.sections:
            for paragraph in section.header.paragraphs:
                yield paragraph
            for paragraph in section.footer.paragraphs:
                yield paragraph

    for para in _iter_all_paragraphs(doc):
        merged = "".join(r.text for r in para.runs)
        if re.search(r"\{\{\s*Name_\s*College\s*\}\}", merged) and "Dean/HOD" in merged:
            for run in para.runs:
                run.text = ""
            if para.runs:
                para.runs[0].text = hod_val

    replace_all(doc, replacements)

    # ── Copy-to section: bold items 1, 2, 5; rebuild item 2 (HOD, College) ───
    copy_idx = next(
        (i for i, p in enumerate(doc.paragraphs) if p.text.strip() == "Copy to:"),
        None,
    )
    if copy_idx is not None:
        paras = doc.paragraphs
        item1 = paras[copy_idx + 1]   # student name
        item2 = paras[copy_idx + 2]   # college HOD line (already set above)
        item5 = paras[copy_idx + 5]   # project guide line
        # item2 now has hod_val; rebuild as "HOD, College" bold
        new_item2 = f"{hod_val}, {college_val}" if hod_val and college_val else (hod_val or college_val)
        _set_para_bold(item2, new_item2)
        _make_para_bold(item1)
        _make_para_bold(item5)

    # Bold the student name (with salute) and the date range in the body paragraph
    salute    = student.get("salute", "").strip()
    name_bold = student.get("name", "").strip()
    from_date = format_date(student.get("start date", ""))
    to_date   = format_date(student.get("end date", ""))
    name_frag = (salute + " " + name_bold).strip() if salute else name_bold
    bold_frags = [f for f in [name_frag, from_date, to_date] if f]
    if bold_frags and name_bold:
        for para in doc.paragraphs:
            if name_bold in para.text and from_date in para.text:
                bold_fragments_in_paragraph(para, bold_frags)
                break

    adjust_font_sizes(doc, target_pt=10)

    # Restore header paragraph sizes (first 4 paragraphs) which were shrunk above
    header_sizes = [Pt(15), Pt(14), Pt(13), Pt(13)]
    for para, size in zip(doc.paragraphs[:4], header_sizes):
        for run in para.runs:
            run.font.size = size

    name      = student.get("name", f"Row_{student['_row']}").strip()
    safe_name = "".join(c if c.isalnum() or c in " _-" else "_" for c in name).strip()
    out_path  = os.path.join(output_folder, f"GA_Letter_{safe_name}.docx")
    doc.save(out_path)
    return out_path


def strip_images(doc):
    """Remove only the signature drawing, keeping the stamp."""
    from docx.oxml.ns import qn
    containers = list(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                containers.extend(cell.paragraphs)
    for para in containers:
        for drawing in para._p.findall('.//' + qn('w:drawing')):
            docPr = drawing.find('.//' + qn('wp:docPr'))
            name  = (docPr.get('name', '') if docPr is not None else '').lower()
            desc  = (docPr.get('descr', '') if docPr is not None else '').lower()
            # Remove only drawings identified as a signature; keep stamp
            if 'sign' in name or 'sign' in desc:
                drawing.getparent().remove(drawing)


def print_table(students):
    print(f"\n  {'Row':<5} {'Name':<25} {'Guide Name':<22} {'Start Date':<18} {'End Date'}")
    print(f"  {'─'*5} {'─'*25} {'─'*22} {'─'*18} {'─'*16}")
    for s in students:
        print(
            f"  {s['_row']:<5} "
            f"{s.get('name', '')[:24]:<25} "
            f"{s.get('guide name', '')[:21]:<22} "
            f"{s.get('start date', '')[:17]:<18} "
            f"{s.get('end date', '')[:15]}"
        )
    print()


def parse_selection(raw, valid_rows):
    raw = (raw or "").strip()

    if not raw:
        return []

    if raw.lower() == "all":
        return sorted(valid_rows)

    selected = []
    for part in raw.replace(" ", "").split(","):
        try:
            r = int(part)
            if r in valid_rows:
                selected.append(r)
            else:
                print(f"  Row {r} not found — skipped.")
        except ValueError:
            print(f"  '{part}' is not a valid number — skipped.")

    return selected


def get_selection(students):
    valid = {s["_row"] for s in students}
    print(f"  Valid rows: {min(valid)} – {max(valid)}")
    print("  Examples:  2        → single letter")
    print("             2,4,5    → multiple letters")
    print("             all      → every student\n")
    raw = input("  Enter row numbers: ").strip()

    return parse_selection(raw, valid)


def main():
    parser = argparse.ArgumentParser(description="Generate GA letters from Google Sheet or local Excel source")
    parser.add_argument("--source", default=EXCEL_FILE, help="Google Sheet URL or local .xlsx path")
    parser.add_argument("--template", default=TEMPLATE_FILE, help="Template .docx file name or path")
    parser.add_argument("--output", default=OUTPUT_FOLDER, help="Output folder path")
    parser.add_argument("--student-json", default="", help="JSON object for a single student record")
    parser.add_argument("--rows", default="", help="Rows to generate, e.g. '2,4,7' or 'all'")
    parser.add_argument("--email", default="", help="Generate letter for a specific student email")
    parser.add_argument("--all", action="store_true", help="Generate letters for all rows")
    parser.add_argument("--non-interactive", action="store_true", help="Skip prompt and use --rows/--all")
    parser.add_argument("--no-table", action="store_true", help="Skip preview table output")
    parser.add_argument("--print", action="store_true", dest="send_to_printer", help="Send to default printer instead of saving to output")
    parser.add_argument("--no-sign", action="store_true", dest="no_sign", help="Strip signature image from the generated letter")
    args = parser.parse_args()

    print("\n" + "=" * 54)
    print("   NRSC GA Letter Generator")
    print("=" * 54)

    script_dir = os.path.dirname(os.path.abspath(__file__))
    template_path = args.template if os.path.isabs(args.template) else os.path.join(script_dir, args.template)

    output_folder = args.output if os.path.isabs(args.output) else os.path.join(script_dir, args.output)

    if not os.path.exists(template_path):
        print(f"\nTemplate not found: {template_path}")
        print(f"Make sure '{TEMPLATE_FILE}' is in the same folder as this script.")
        sys.exit(1)

    os.makedirs(output_folder, exist_ok=True)

    if args.student_json:
        try:
            parsed = json.loads(args.student_json)
        except json.JSONDecodeError as exc:
            print(f"Invalid --student-json payload: {exc}")
            sys.exit(1)

        if not isinstance(parsed, dict):
            print("--student-json must be a single JSON object")
            sys.exit(1)

        student = {
            "_row": 1,
            "salute": str(parsed.get("salute", "")).strip(),
            "name": str(parsed.get("name", "")).strip(),
            "college name": str(parsed.get("college") or parsed.get("college_name") or parsed.get("college name") or "").strip(),
            "district": str(parsed.get("college_district") or parsed.get("district") or "").strip(),
            "college dean hod": str(parsed.get("hod_name") or parsed.get("college_dean_hod") or parsed.get("college dean hod") or "").strip(),
            "guide name": str(parsed.get("guide_name") or parsed.get("guide name") or "").strip(),
            "guide area": str(parsed.get("guide_area") or parsed.get("guide area") or "").strip(),
            "course": str(parsed.get("course") or parsed.get("qualification") or "").strip(),
            "start date": str(parsed.get("start_date") or parsed.get("start date") or "").strip(),
            "end date": str(parsed.get("end_date") or parsed.get("end date") or "").strip(),
            "allotment date": str(parsed.get("allotment_date") or parsed.get("allotment date") or "").strip(),
            "email": str(parsed.get("email", "")).strip(),
            "month": str(parsed.get("month", "")).strip(),
            "sl. no.": str(parsed.get("sl_no") or parsed.get("sl. no.") or parsed.get("student_uid") or "").strip(),
            "guide reporting officer": str(parsed.get("guide_reporting_officer") or parsed.get("guide reporting officer") or "").strip(),
            "dd": str(parsed.get("dd", "")).strip(),
        }
        students = [student]
    else:
        students = load_students(args.source)

    print(f"  {len(students)} student record(s) found:\n")

    if not args.no_table:
        print_table(students)

    valid_rows = {s["_row"] for s in students}

    if args.student_json:
        selected_rows = [1]
    elif args.email:
        requested_email = args.email.strip().lower()
        selected_rows = [s["_row"] for s in students if s.get("email", "").strip().lower() == requested_email]
    elif args.all:
        selected_rows = sorted(valid_rows)
    elif args.rows:
        selected_rows = parse_selection(args.rows, valid_rows)
    elif args.non_interactive:
        selected_rows = sorted(valid_rows)
    else:
        selected_rows = get_selection(students)

    if not selected_rows:
        if args.email:
            print(f"\nNo student found for email: {args.email}")
            sys.exit(1)
        print("\nNo valid rows selected. Exiting.")
        sys.exit(0)

    row_map = {s["_row"]: s for s in students}
    print(f"\nGenerating {len(selected_rows)} letter(s)...\n")

    success, failed = [], []
    for row_num in selected_rows:
        student = row_map[row_num]
        try:
            out_path = generate_letter(template_path, student, output_folder)
            # Always strip signature for the open-in-Word path
            if args.send_to_printer or args.no_sign:
                from docx import Document as _Doc
                _d = _Doc(out_path)
                strip_images(_d)
                _d.save(out_path)
            print(f"  [{row_num}]  {student.get('name', ''):<28} ->  {os.path.basename(out_path)}")
            if args.send_to_printer:
                print_docx(out_path)
                print(f"         Sent to printer.")
            success.append(out_path)
        except Exception as e:
            print(f"  [{row_num}]  Error: {e}")
            import traceback; traceback.print_exc()
            failed.append(row_num)

    print(f"\n{'='*54}")
    print(f"  {len(success)} letter(s) saved to  ->  {output_folder}")
    if args.send_to_printer:
        print(f"  {len(success)} letter(s) sent to printer")
    if failed:
        print(f"  Failed rows: {failed}")
    print("=" * 54 + "\n")
    # When a single letter is generated (web API path), output the exact file path
    result_output = success[0] if len(success) == 1 else output_folder
    print(f"RESULT: success={len(success)} failed={len(failed)} output={result_output}")
    if failed:
        sys.exit(1)


if __name__ == "__main__":
    main()