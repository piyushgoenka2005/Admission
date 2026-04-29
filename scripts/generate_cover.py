"""
Generate a Project Closure Certificate (.docx) for a student using the
template file  'project closing certificate.docx'.

Usage (from the web app):
    python generate_closure.py --student-json '{"name":"...","start_date":"...","end_date":"..."}'

The certificate is saved into the 'output' folder.
"""

import os
import sys
import json
import argparse
import datetime
import tempfile
import time
from zoneinfo import ZoneInfo

from docx import Document
from docx.shared import Pt

TEMPLATE_FILE = "../sample/cover.docx"
OUTPUT_FOLDER = "output"
HYDERABAD_TZ = ZoneInfo("Asia/Kolkata")


VIRTUAL_PRINTERS = {"microsoft print to pdf", "microsoft xps document writer",
                     "onenote", "onenote (desktop)", "onenote (desktop) - protected",
                     "fax", "send to onenote"}


def find_physical_printer():
    """
    Return the printer to use.
    Priority:
      1. NRSC_PRINTER environment variable (set this to the exact LAN printer name)
      2. First non-virtual printer found on the system
    """
    env_printer = os.environ.get("NRSC_PRINTER", "").strip()
    if env_printer:
        return env_printer
    import subprocess
    result = subprocess.run(
        ["powershell", "-NoProfile", "-Command",
         "Get-Printer | Select-Object -ExpandProperty Name"],
        capture_output=True, text=True
    )
    for line in result.stdout.strip().splitlines():
        name = line.strip()
        if name and name.lower() not in VIRTUAL_PRINTERS:
            return name
    return None


def print_docx(filepath):
    """Open .docx in Word visibly for review — user can edit and print manually."""
    import pythoncom
    import win32com.client

    abs_path = os.path.abspath(filepath)
    pythoncom.CoInitialize()
    try:
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = True
        word.WindowState = 1            # wdWindowStateNormal
        word.Documents.Open(abs_path)
        # Leave Word open — user reviews/prints manually
    except Exception:
        raise
    finally:
        pythoncom.CoUninitialize()


# ─────────────────────────────────────────────────────────────────────────────
# Helpers
# ─────────────────────────────────────────────────────────────────────────────

def format_date(date_str: str) -> str:
    """Parse an ISO-ish date string and return  DD Mon YYYY  (e.g. 01 Feb 2026)."""
    if not date_str:
        return ""
    try:
        raw = str(date_str).strip()
        if raw.lower() in {"undefined", "null"}:
            return ""
        dt = datetime.datetime.fromisoformat(raw.replace(" ", "T").split("T")[0])
        return dt.strftime("%d %b %Y")
    except ValueError:
        return date_str


def format_datetime(date_str: str) -> str:
    """Parse an ISO-ish datetime string and return  DD Mon YYYY, HH:MM AM/PM."""
    if not date_str:
        return ""
    try:
        raw = str(date_str).strip()
        if raw.lower() in {"undefined", "null"}:
            return ""

        dt = datetime.datetime.fromisoformat(raw.replace("Z", "+00:00"))
        if dt.tzinfo is not None:
            dt = dt.astimezone(HYDERABAD_TZ)
        return dt.strftime("%d %b %Y, %I:%M %p")
    except ValueError:
        return format_date(date_str)


def clean_name_text(value: str) -> str:
    text = str(value or "").strip()
    if text.lower() in {"undefined", "null"}:
        return ""
    return text


def calculate_age_from_dob(dob_value, fallback_age=""):
    """Calculate current age from DOB in Hyderabad local date context."""
    if not dob_value:
        return str(fallback_age or "")

    try:
        dob_text = str(dob_value).strip().split("T")[0]
        if dob_text.lower() in {"undefined", "null"}:
            return str(fallback_age or "")

        dob_date = datetime.date.fromisoformat(dob_text)
        today = datetime.datetime.now(HYDERABAD_TZ).date()

        age = today.year - dob_date.year
        if (today.month, today.day) < (dob_date.month, dob_date.day):
            age -= 1

        return str(age if age >= 0 else "")
    except Exception:
        return str(fallback_age or "")


def build_replacements(student: dict) -> dict:
    """Return placeholder -> value mapping for the cover application template."""
    name = clean_name_text(student.get("student_name") or student.get("name", ""))
    if not name:
        first = clean_name_text(student.get("first_name", ""))
        middle = clean_name_text(student.get("middle_name", ""))
        last = clean_name_text(student.get("last_name", ""))
        name = " ".join([part for part in [first, middle, last] if part]).strip()

    # Get start date and end date from new or old keys
    start_date = student.get("present_date") or student.get("start_date", "")
    end_date = student.get("date_of_ending") or student.get("end_date", "")
    
    # Format duration for display
    duration = student.get("duration", "")
    duration_display = duration.replace("_", " ").title() if duration else ""
    
    # Format family members
    family_members = student.get("family_members", [])
    family_members_text = ""
    if family_members and isinstance(family_members, list):
        family_members_text = ", ".join([
            f"{m.get('name', '')} ({m.get('relationship', '')})" 
            for m in family_members 
            if m.get('name')
        ])
    
    # Format entry permissions
    entry_permissions = student.get("entry_permissions", [])
    entry_perms_text = ""
    if entry_permissions and isinstance(entry_permissions, list):
        allowed_locations = [
            p.get('location', '') 
            for p in entry_permissions 
            if p.get('allowed') and p.get('location')
        ]
        entry_perms_text = ", ".join(allowed_locations)
    
    # Photo ID proofs
    photo_id_proof = student.get("photo_id_proof", "")
    
    replacements = {
        # Basic Info
        "{{Student_Name}}":       name,
        "{{Salutation}}":         student.get("salutation", ""),
        "{{Gender}}":             student.get("gender", ""),
        "{{State}}":              student.get("state", ""),
        "{{City}}":               student.get("city", ""),
        
        # Contact
        "{{Phone}}":              student.get("phone_number", "") or student.get("phone", ""),
        "{{Phone_Number}}":       student.get("phone_number", "") or student.get("phone", ""),
        "{{Email}}":              student.get("email", ""),
        "{{email}}":              student.get("email", ""),
        
        # Personal Details
        "{{Date_of_Birth}}":      student.get("date_of_birth", ""),
        "{{Date_Of_Birth}}":      format_date(student.get("date_of_birth", "")),
        "{{Age}}":                calculate_age_from_dob(student.get("date_of_birth", ""), student.get("age", "")),
        
        # Family Details
        "{{Fathers_Name}}":       student.get("fathers_name", ""),
        "{{Mothers_Name}}":       student.get("mothers_name", ""),
        
        # Educational Details
        "{{College_Name}}":       student.get("college_name", "") or student.get("university_name", ""),
        "{{College_Phone}}":      student.get("college_phone_number", ""),
        "{{College_Phone_Number}}": student.get("college_phone_number", ""),
        "{{College_District}}":   student.get("college_district", ""),
        "{{College_Address}}":    student.get("college_address", ""),
        "{{University}}":         student.get("university_name", ""),
        "{{University_Name}}":    student.get("university_name", "") or student.get("college_name", ""),
        "{{Qualification}}":      student.get("qualification", ""),
        "{{Degree}}":             student.get("qualification", ""),
        "{{HOD_Name}}":           student.get("hod_name", ""),
        "{{HOD_Email}}":          student.get("hod_email", ""),
        "{{hod_email}}":          student.get("hod_email", ""),
        
        # Address Details
        "{{Permanent_Address}}":  student.get("permanent_address", ""),
        "{{Present_Address}}":    student.get("present_address", "") or student.get("residential_address", ""),
        
        # Internship Details
        "{{Duration}}":           duration_display,
        "{{Contact_No}}":         student.get("phone_number", "") or student.get("phone", ""),
        "{{Date_and_time}}":      format_datetime(student.get("reported_at") or student.get("reporting_date") or start_date),
        "{{Reporting_Date}}":     format_datetime(student.get("reported_at") or student.get("reporting_date")),
        "{{Reporting _Date_&_Time}}": format_datetime(student.get("reported_at") or student.get("reporting_date") or start_date),
        "{{Start_Date}}":         format_date(start_date),
        "{{End_Date}}":           format_date(end_date),
        "{{Present_Date}}":       format_date(start_date),
        "{{Date_Of_Ending}}":     format_date(end_date),
        "{{Reporting_Division}}": student.get("reporting_division", ""),
        "{{Entry_Permissions}}":  entry_perms_text,
        
        # Guide Details
        "{{Guide_Name}}":         student.get("assigned_scientist", "") or student.get("guide_name", ""),
        "{{Guide_Email}}":        student.get("guide_email", ""),
        "{{Guide_email}}":        student.get("guide_email", ""),
        "{{Guide_Phone}}":        student.get("guide_phone", ""),
        "{{Guide_D&T}}":          format_datetime(student.get("guide_assigned_at", "")),
        "{{Guide_Assigned_At}}":  format_datetime(student.get("guide_assigned_at", "")),
    }

    # Sanitize all values to remove newlines that could push layout down
    for k, v in replacements.items():
        if isinstance(v, str):
            replacements[k] = v.replace("\n", ", ").replace("\r", "").strip()
            
    return replacements


def replace_in_paragraph(para, replacements):
    """
    Replace every placeholder in *para*, preserving run formatting.
    """
    for ph, val in replacements.items():
        if not para.runs:
            continue
        
        # Pass 1: replace within individual runs (fast and preserves per-run formatting)
        for run in para.runs:
            if ph in run.text:
                run.text = run.text.replace(ph, val)
                
        # Pass 2: handle placeholders split across multiple runs
        full_text = "".join(r.text for r in para.runs)
        while ph in full_text:
            start_idx = full_text.find(ph)
            
            current_idx = 0
            start_run_idx = -1
            start_char_idx = -1
            for i, run in enumerate(para.runs):
                if current_idx + len(run.text) > start_idx:
                    start_run_idx = i
                    start_char_idx = start_idx - current_idx
                    break
                current_idx += len(run.text)
                
            end_idx = start_idx + len(ph)
            current_idx = 0
            end_run_idx = -1
            end_char_idx = -1
            for i, run in enumerate(para.runs):
                if current_idx + len(run.text) >= end_idx:
                    end_run_idx = i
                    end_char_idx = end_idx - current_idx
                    break
                current_idx += len(run.text)
                
            before_text = para.runs[start_run_idx].text[:start_char_idx]
            after_text = para.runs[end_run_idx].text[end_char_idx:]
            
            para.runs[start_run_idx].text = before_text + val
            if start_run_idx == end_run_idx:
                para.runs[start_run_idx].text += after_text
            else:
                para.runs[end_run_idx].text = after_text
                for i in range(start_run_idx + 1, end_run_idx):
                    para.runs[i].text = ""
                    
            full_text = "".join(r.text for r in para.runs)


def replace_all(doc, replacements):
    """Apply replacements to every paragraph, table cell, header, footer, and textboxes."""
    from docx.text.paragraph import Paragraph
    
    for para in doc.paragraphs:
        replace_in_paragraph(para, replacements)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    replace_in_paragraph(para, replacements)

    for section in doc.sections:
        for para in section.header.paragraphs:
            replace_in_paragraph(para, replacements)
        for para in section.footer.paragraphs:
            replace_in_paragraph(para, replacements)
            
    # Replace inside textboxes (w:txbxContent)
    def replace_in_element_tree(element):
        for node in element.iter():
            if node.tag.endswith('}txbxContent'):
                for p in node.findall('.//w:p', namespaces=node.nsmap):
                    para = Paragraph(p, doc)
                    replace_in_paragraph(para, replacements)
                    
    replace_in_element_tree(doc.element.body)
    for section in doc.sections:
        replace_in_element_tree(section.header._element)
        replace_in_element_tree(section.footer._element)


# ─────────────────────────────────────────────────────────────────────────────
# Certificate generation
# ─────────────────────────────────────────────────────────────────────────────




def generate_cover_document(template_path: str, student: dict, output_folder: str) -> str:
    """Open the template, replace placeholders, save to output folder."""
    doc          = Document(template_path)
    replacements = build_replacements(student)

    replace_all(doc, replacements)

    name = clean_name_text(student.get("student_name") or student.get("name", ""))
    if not name:
        first = clean_name_text(student.get("first_name", ""))
        middle = clean_name_text(student.get("middle_name", ""))
        last = clean_name_text(student.get("last_name", ""))
        name = " ".join([part for part in [first, middle, last] if part]).strip()
    if not name:
        name = "unknown"
        
    safe_name = "".join(c if c.isalnum() or c in " _-" else "_" for c in name).strip()
    os.makedirs(output_folder, exist_ok=True)
    out_path  = os.path.join(output_folder, f"Application_Form_{safe_name}.docx")
    doc.save(out_path)
    return out_path


def main():
    parser = argparse.ArgumentParser(description="Generate Application Cover Document")
    parser.add_argument("--student-json", required=True, help="JSON object with student data")
    parser.add_argument("--template", default=TEMPLATE_FILE, help="Template .docx file")
    parser.add_argument("--output", default=OUTPUT_FOLDER, help="Output folder path")
    parser.add_argument("--print", action="store_true", dest="send_to_printer", help="Send to default printer instead of saving to output")
    args = parser.parse_args()

    script_dir = os.path.dirname(os.path.abspath(__file__))
    template_path = args.template if os.path.isabs(args.template) else os.path.join(script_dir, args.template)

    output_folder = args.output if os.path.isabs(args.output) else os.path.join(script_dir, args.output)

    if not os.path.exists(template_path):
        print(f"\nTemplate not found: {template_path}")
        print(f"Make sure '{TEMPLATE_FILE}' is in the same folder as this script.")
        sys.exit(1)

    try:
        student = json.loads(args.student_json)
    except json.JSONDecodeError as exc:
        print(f"Invalid --student-json payload: {exc}")
        sys.exit(1)

    if not isinstance(student, dict):
        print("--student-json must be a single JSON object")
        sys.exit(1)

    print("\n" + "=" * 54)
    print("   Application Cover Generator")
    print("=" * 54)

    try:
        out_path = generate_cover_document(template_path, student, output_folder)
        if args.send_to_printer:
            print_docx(out_path)
            print(f"\n  Certificate sent to printer -> {os.path.basename(out_path)}")
            print(f"\n{'='*54}")
            print(f"  1 certificate sent to printer")
            # clean up temp folder
            try:
                os.rmdir(output_folder)
            except OSError:
                pass
        else:
            print(f"\n  Certificate saved -> {os.path.basename(out_path)}")
            print(f"\n{'='*54}")
            print(f"  1 certificate saved to  ->  {output_folder}")
        print("=" * 54 + "\n")
        print(f"RESULT: success=1 failed=0 output={output_folder}")
    except Exception as e:
        print(f"\n  Error: {e}")
        import traceback
        traceback.print_exc()
        print(f"RESULT: success=0 failed=1 output={output_folder}")
        sys.exit(1)


if __name__ == "__main__":
    main()
