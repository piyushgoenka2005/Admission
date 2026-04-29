"""
Generate a memo requesting ACS-Biometric & VM Login credentials for
Project Students, matching the format from ID-card.docx.

Usage:
    python generate_idcard_sheet.py --students-json '[{...}, {...}]' [--print]
"""

import os
import sys
import json
import argparse
import tempfile
import time
import datetime

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

VIRTUAL_PRINTERS = {"microsoft print to pdf", "microsoft xps document writer",
                     "onenote", "onenote (desktop)", "onenote (desktop) - protected",
                     "fax", "send to onenote"}


def find_physical_printer():
    """
    Return the printer to use.
    Priority:
      1. NRSC_PRINTER environment variable (set this to the exact LAN printer name)
      2. First non-virtual printer found on the system
    """
    env_printer = os.environ.get("NRSC_PRINTER", "").strip()
    if env_printer:
        return env_printer
    import subprocess
    result = subprocess.run(
        ["powershell", "-NoProfile", "-Command",
         "Get-Printer | Select-Object -ExpandProperty Name"],
        capture_output=True, text=True
    )
    for line in result.stdout.strip().splitlines():
        name = line.strip()
        if name and name.lower() not in VIRTUAL_PRINTERS:
            return name
    return None


def print_docx(filepath):
    """Open .docx in Word visibly for review — user can edit and print manually."""
    import pythoncom
    import win32com.client

    abs_path = os.path.abspath(filepath)
    pythoncom.CoInitialize()
    try:
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = True
        word.WindowState = 1            # wdWindowStateNormal
        word.Documents.Open(abs_path)
        # Leave Word open — user reviews/prints manually
    except Exception:
        raise
    finally:
        pythoncom.CoUninitialize()


def format_date_short(date_str):
    """Parse ISO date and return DD-Mon-YY (e.g. 05-Feb-26)."""
    if not date_str:
        return ""
    for fmt in ("%Y-%m-%d", "%d-%m-%Y", "%Y-%m-%dT%H:%M:%S.%fZ"):
        try:
            d = datetime.datetime.strptime(date_str.strip(), fmt)
            return d.strftime("%d-%b-%y")
        except ValueError:
            continue
    return date_str


def format_date_long(date_str):
    """Parse ISO date and return readable date like '3rd Feb 2026'."""
    if not date_str:
        return ""
    for fmt in ("%Y-%m-%d", "%d-%m-%Y"):
        try:
            d = datetime.datetime.strptime(date_str.strip(), fmt)
            day = d.day
            if 11 <= day <= 13:
                suffix = "th"
            else:
                suffix = {1: "st", 2: "nd", 3: "rd"}.get(day % 10, "th")
            return f"{day}{suffix} {d.strftime('%b %Y')}"
        except ValueError:
            continue
    return date_str


def add_para(doc, text="", bold=False, size=11, name="Calibri",
             align=WD_ALIGN_PARAGRAPH.LEFT, space_after=Pt(4)):
    """Add a paragraph with a single run."""
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = space_after
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = name
    return p, run


def set_cell(cell, text, bold=False, size=10, align=WD_ALIGN_PARAGRAPH.LEFT,
             font_name="Calibri"):
    """Set cell text with formatting."""
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run(str(text))
    run.font.size = Pt(size)
    run.font.name = font_name
    run.bold = bold
    return run


def shade_cell(cell, fill_color):
    """Apply background shading to a cell."""
    tc_pr = cell._element.get_or_add_tcPr()
    shd = tc_pr.makeelement(qn("w:shd"), {
        qn("w:val"): "clear",
        qn("w:color"): "auto",
        qn("w:fill"): fill_color,
    })
    tc_pr.append(shd)


def generate_datasheet(students, do_print=False, output_dir="output"):
    """Populate the ID-card.docx template with the selected students."""
    script_dir    = os.path.dirname(os.path.abspath(__file__))
    template_path = os.path.join(script_dir, "ID-card.docx")

    doc = Document(template_path)

    # ── Replace {{Date}} placeholder with today's date ───────────────────
    today = datetime.date.today()
    day   = today.day
    if 11 <= day <= 13:
        suffix = "th"
    else:
        suffix = {1: "st", 2: "nd", 3: "rd"}.get(day % 10, "th")
    date_str = f"{day}{suffix} {today.strftime('%b %Y')}"

    for para in doc.paragraphs:
        # Per-run pass
        for run in para.runs:
            if "{{Date}}" in run.text:
                run.text = run.text.replace("{{Date}}", date_str)
        # Merged-run pass (handles placeholder split across runs)
        full = "".join(r.text for r in para.runs)
        if "{{Date}}" in full:
            full = full.replace("{{Date}}", date_str)
            if para.runs:
                para.runs[0].text = full
                for r in para.runs[1:]:
                    r.text = ""

    # ── Clear sample data rows (keep header row 0) ───────────────────────
    table = doc.tables[0]
    while len(table.rows) > 1:
        table._element.remove(table.rows[-1]._tr)

    # ── Add a row per selected student ───────────────────────────────────
    col_widths = [Cm(1.3), Cm(1.6), Cm(5.2), Cm(2.5), Cm(2.8), Cm(2.8)]
    for i, s in enumerate(students, 1):
        row = table.add_row()
        values = [
            f"{i}.",
            s.get("salute", ""),
            s.get("name", ""),
            format_date_short(s.get("start_date", "")),
            format_date_short(s.get("end_date", "")),
            s.get("location", "").strip() or "Jeedimetla",
        ]
        for idx, val in enumerate(values):
            align = WD_ALIGN_PARAGRAPH.CENTER if idx in (0, 3, 4) else WD_ALIGN_PARAGRAPH.LEFT
            set_cell(row.cells[idx], val, size=10, align=align,
                     font_name="Times New Roman")
            row.cells[idx].width = col_widths[idx]

    # ── Save / Print ─────────────────────────────────────────────────────
    os.makedirs(output_dir, exist_ok=True)
    out_path = os.path.join(
        output_dir,
        f"ACS_Biometric_Request_{today.strftime('%Y%m%d')}.docx"
    )
    doc.save(out_path)
    if do_print:
        print_docx(out_path)
        print(f"RESULT: success=1 failed=0 output={out_path} printed=1")
    else:
        print(f"RESULT: success=1 failed=0 output={out_path}")


def main():
    parser = argparse.ArgumentParser(description="Generate ID Card Data Sheet")
    parser.add_argument("--students-json", required=True,
                        help="JSON array of student objects")
    parser.add_argument("--output", default="output", help="Output folder path")
    parser.add_argument("--print", dest="do_print", action="store_true",
                        help="Print directly instead of saving")
    args = parser.parse_args()

    students = json.loads(args.students_json)
    if not students:
        print("No students provided", file=sys.stderr)
        sys.exit(1)

    generate_datasheet(students, do_print=args.do_print, output_dir=args.output)


if __name__ == "__main__":
    main()
